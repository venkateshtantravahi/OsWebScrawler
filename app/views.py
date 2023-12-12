import os, json, time
from threading import Thread, Event
from flask import render_template, request, url_for, redirect, send_from_directory, jsonify, session, send_file
from flask_login import login_user, logout_user, current_user, login_required
from flask_wtf.csrf import generate_csrf
from jinja2 import TemplateNotFound
import pandas as pd
from docx import Document
from sqlalchemy import or_

from . import app, lm, bc
from app.forms import LoginForm, RegisterForm
from app.models import Users, Data, CrawledData, db
from app.services.queue_service import RequestQueue
from app.services.crawler import WebCrawler, MEDIA_DIR
from app.utils.text_sanitizer import sanitize_text


request_queue = RequestQueue()

stop_event = Event()

results = []

user_id = None

def crawl_worker():
    """
    Worker function for handling crawl tasks.
    
    Continuously checks the request queue for new crawl tasks, processes them, 
    and stores the results in the database.
    """
    while not stop_event.is_set():
        if not request_queue.is_empty():
            task = request_queue.get_request()
            if task:
                try:
                    task = json.loads(task)
                    print(f"Processing URL {task['url']} for user {task['user_id']}")
                    # Process task...
                    user_id = task['user_id']
                    url = task['url'] 
                    print(f'{url}: {user_id}')
                    web_crawler = WebCrawler(url=task['url'])
                    if url:
                        try:
                            result = web_crawler.crawl(task['url'])
                            if result:
                                print(f"retrieved results waiting to save in db {task['url']} {task['user_id']}")
                                # Create an instance of CrawledData and populate its fields
                                crawled_data = CrawledData( 
                                    user_id=task['user_id'],
                                    url=result['url'],
                                    title=result['title'],
                                    content=result['content'],
                                    file_path=result['file_path'] if result['file_path'] else '',
                                    content_type=result['content-type'],
                                    links=','.join(result['links']) if result['links'] else ''  # Assuming the links column is a comma-separated string of links
                                )
                                print("crawl object created!")
                                crawled_data.save()
                                # flash("URL has been scraped", 'success')
                                results.append(result)
                        except Exception as e:
                            print(f"An error occured: {e}")
                except json.JSONDecodeError as e:
                    print(f"JSON decode error: {e}")
                    continue  # Skip to the next iteration if decoding fails


@lm.user_loader
def load_user(user_id):
    """
    Flask-Login user_loader callback.
    
    Loads a user from the database using the given user_id.
    
    :param user_id: int - The user's identifier.
    :return: User object or None if the user doesn't exist.
    """
    return Users.query.get(int(user_id))


@app.route('/logout')
@login_required
def logout():
    """
    Route to log out the current user.
    
    Ends the user's session and redirects to the index page.
    """
    logout_user()
    return redirect(url_for('index'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    """
    Route to handle user registration.
    
    GET: Renders the registration form.
    POST: Processes the registration form and creates a new user.
    :return: Redirection or JSON response with registration status.
    """
    if request.method == 'POST':
        form = RegisterForm(request.form)

        if form.validate_on_submit():
            #Fitler data for existing users
            user = Users.query.filter_by(user=form.username.data).first()
            user_by_email = Users.query.filter_by(email=form.email.data).first()
            if user or user_by_email is None:
                pwd_hash = bc.generate_password_hash(form.password.data)
                user = Users(firstname=form.firstname.data,
                             lastname=form.lastname.data,
                             user=form.username.data,
                             password=pwd_hash,
                             email=form.email.data)
                user.save()
                return redirect(url_for('index')), 201

            elif user or user_by_email:
                # flash("User already exists", 'error')
                return jsonify({"message": "User already exists."}), 409
        else:
            return jsonify({"errors": form.errors}), 400
        
    else:
        form = RegisterForm()
        return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    """
    Route to handle user login.
    
    GET: Renders the login form.
    POST: Authenticates the user and initiates a session.
    :return: Redirection to the index page or rendering of login form with error.
    """
    if request.method == 'POST':
        form = LoginForm(request.form)

        if form.validate_on_submit():
            username = Users.query.filter_by(user=form.username.data).first()
            if username and bc.check_password_hash(username.password, form.password.data):
                login_user(username, remember=form.remember.data)
                user_id = username.id
                # user_id = current_user.get_id()
                return redirect(url_for('index'))
            
            else:
                return render_template('login.html', form=form, error="Invalid credentials.")
        else:
            # flash("form errors: {}".format(form.errors), 'error')
            return jsonify({"errors": form.errors}), 400
    else:
        form = LoginForm()
        # session['csrf_token'] = generate_csrf()
        return render_template('login.html', form=form)

@app.route('/enqueue', methods=['POST'])
@login_required
def enqueue_url():
    """
    Route to enqueue a URL for crawling.
    
    Accepts a URL and optional content type, and adds it to the crawl queue.
    :return: JSON response indicating the URL was added to the queue.
    """
    data = request.json
    url = data.get('url')
    content_type = data.get('content_type', 'other')  # Default to 'other' if not specified
    user_id = current_user.get_id()  # Get user ID from the current_user
    request_queue.add_request(user_id=user_id, url=url, content_type=content_type)
    # flash("URL {} added to request queue".format(url), 'success')
    return jsonify({"message": "URL added to queue", "url": url}), 200

# Starting multiple worker threads for scalability
num_worker_threads = 5  # Adjust this number based on your needs and server capacity
workers = []
def thread_with_context():
    """
    Wrapper function to run the crawl worker in a Flask application context.

    No input parameters.
    Outputs: None. The function works by side-effects, running the crawl worker.
    """
    with app.app_context():
        crawl_worker()

for _ in range(num_worker_threads):
    thread = Thread(target=thread_with_context) 
    thread.daemon = True
    thread.start() 
    workers.append(thread)

# @app.route('/start_crawling', methods=['GET'])
# @login_required
# def start_crawling():
#     # This endpoint could trigger the crawl process if not already running
#     # For now, it just returns the status
#     return jsonify({"message": "Crawler is running"}), 200

@app.route('/download-my-data')
@login_required
def download_my_data():
    """
    Route to allow users to download their data.
    
    Fetches user-specific crawled data from the database and provides a Word document for download.
    :return: Word document with user data or a 404 error if no data is available.
    """
    user_id = current_user.get_id()  # Get the current logged-in user's ID
    data = CrawledData.query.filter_by(user_id=user_id).all()  # Fetch user-specific data

    if not data:
        app.logger.info("No data available")
        # flash("No data available for user {}".format(user_id), 'warning')
        app.logger.log(f"No data available for user {user_id}")
        return "No data available", 404

    # Create a Word document
    doc = Document()
    for item in data:
        url = sanitize_text(item.url)
        title = sanitize_text(item.title)
        content = sanitize_text(item.content)
        file_path = sanitize_text(item.file_path)
        links = ', '.join(sanitize_text(link) for link in (item.links or '').split(','))

        doc.add_heading('URL: ' + url, level=1)
        doc.add_paragraph('Title: ' + title)
        doc.add_paragraph('Content: ' + content)
        doc.add_paragraph('File Path: ' + (file_path or 'N/A'))
        doc.add_paragraph('Links: ' + ', '.join(links))

    # Save the document to a temporary file
    filename = f"user_data_{user_id}.docx"
    file_path = os.path.join( MEDIA_DIR, filename)
    print(file_path)
    doc.save(file_path)

    return send_file(file_path, as_attachment=True) #, attachment_filename=filename

@app.route('/stop_crawling', methods=['POST'])
@login_required
def stop_crawling():
    """
    Stops the web crawling process. This route sets a stop event and joins all worker threads to ensure
    the crawling process is cleanly terminated.

    This route requires user authentication.

    Method: POST
    URL: /stop_crawling

    Returns:
        JSON response indicating the crawling process has been stopped.
    """
    # This endpoint will stop the crawling process
    stop_event.set()
    for worker in workers:
        worker.join()  # Wait for threads to finish
    # flash("Crawler has been stopped", 'success')
    app.logger.warn("Crawler has been stopped")
    return jsonify({"message": "Crawler has been stopped"}), 200

@app.route('/crawler_control') 
@login_required
def crawler_control():
     """
    Presents the crawler control interface to the user. This route renders a template allowing 
    authenticated users to control the web crawling process.

    This route requires user authentication.

    Method: GET
    URL: /crawler_control

    Returns:
        Rendered HTML template for crawler control.
    """
     return render_template('crawl_data.html')

@app.route('/', defaults={'path': 'index'}, methods=['GET', 'POST'])
@app.route('/<path>', methods=['GET', 'POST'])
def index(path):
    """
    The main index route for the application.
    
    If the user is authenticated, it will render search results based on a query or display 
    all user-specific crawled data.
    
    If the user is not authenticated, it redirects to the login page.
    
    In case of template not found, it returns a 404 error page, and for other exceptions, 
    it returns a 500 error page.
    """
    try:
        if not current_user.is_authenticated:
            # Redirect to the login page
            return redirect(url_for('login'))

        user_id = current_user.get_id()  # Get the current user's ID
        db.session.expire_all()
        query = request.args.get('search')  # Get the search term from the URL query parameters

        if query:
            search_term = f"%{query}%"  # Format the search term for partial matching
            # Query for a partial match in either the title or the content_type
            results = CrawledData.query.filter(
                CrawledData.user_id == user_id,
                or_(
                    CrawledData.title.ilike(search_term),
                    CrawledData.content_type.ilike(search_term)
                )
            ).all()
        else:
            # For a GET request or POST without a query, display all user-specific data
            results = CrawledData.query.filter_by(user_id=user_id).all()

        # Render the index page with the search results or all user-specific data
        return render_template('index.html', results=results, query=query)

    except TemplateNotFound:
        # If the template is not found, render a 404 error page
        return render_template('page-404.html'), 404
    except Exception as e:
        # Log the exception and render a 500 error page
        app.logger.error(f'Error: {e}')
        return render_template('page-500.html'), 500


@app.route('/sitemap.xml')
def sitemap():
    """
    Provides the sitemap.xml file. This route serves the sitemap.xml from the static directory of the 
    application, typically used by search engines for crawling.

    Method: GET
    URL: /sitemap.xml

    Returns:
        sitemap.xml file from the static directory.
    """
    return send_from_directory(os.path.join(app.root_path, 'static'), 'sitemap.xml')




